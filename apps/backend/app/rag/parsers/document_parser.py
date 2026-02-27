from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass(slots=True)
class ParsedSection:
    text: str
    source_page: int | None = None


@dataclass(slots=True)
class ParsedDocumentContent:
    sections: list[ParsedSection]
    metadata: dict


def parse_document_content(
    path: Path,
    *,
    max_pdf_pages: int = 1000,
    pdf_ocr_enabled: bool = False,
) -> ParsedDocumentContent:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return ParsedDocumentContent(
                    sections=[ParsedSection(text=path.read_text(encoding=enc), source_page=None)],
                    metadata={"encoding": enc},
                )
            except UnicodeDecodeError:
                continue
        return ParsedDocumentContent(
            sections=[ParsedSection(text=path.read_text(errors="ignore"), source_page=None)],
            metadata={"encoding": "ignore"},
        )
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        if page_count > max_pdf_pages:
            raise ValueError(f"PDF has {page_count} pages which exceeds the limit of {max_pdf_pages} pages")
        sections: list[ParsedSection] = []
        empty_pages = 0
        for idx, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                empty_pages += 1
            sections.append(ParsedSection(text=text, source_page=idx))
        if page_count > 0 and not pdf_ocr_enabled and empty_pages / page_count >= 0.8:
            raise ValueError(
                "PDF appears to be mostly image/scanned pages and OCR is disabled. "
                "Please upload a text-based PDF or enable OCR support."
            )
        return ParsedDocumentContent(
            sections=sections,
            metadata={
                "pages": page_count,
                "empty_pages": empty_pages,
                "pdf_ocr_enabled": pdf_ocr_enabled,
            },
        )
    if suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        return ParsedDocumentContent(
            sections=[ParsedSection(text=text, source_page=None)],
            metadata={"paragraphs": len(doc.paragraphs)},
        )
    raise ValueError(f"Unsupported document type: {suffix}")


def parse_document_text(path: Path) -> tuple[str, dict]:
    parsed = parse_document_content(path)
    text = "\n\n".join(section.text for section in parsed.sections)
    return text, parsed.metadata
